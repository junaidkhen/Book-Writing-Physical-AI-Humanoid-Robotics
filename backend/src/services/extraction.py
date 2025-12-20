"""Text extraction service for the document ingestion system.

Extracts text content from various document formats (PDF, DOCX, HTML, TXT) with
format-specific extractors per research.md Decision 1. Includes HTML sanitization
to remove scripts and styles. Also extracts document metadata per US4 requirements.
"""
import os
from datetime import datetime
from typing import Optional, Dict, Any
import PyPDF2
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
import requests
from urllib.parse import urljoin, urlparse

from ..utils.logging_config import get_logger


logger = get_logger(__name__)


class TextExtractionError(Exception):
    """Custom exception for text extraction errors."""
    pass


class TextExtractor:
    """Text extraction service with format-specific extractors."""

    @staticmethod
    def extract_metadata_from_pdf(file_path: str) -> Dict[str, Any]:
        """Extract metadata from PDF file.

        Args:
            file_path: Path to PDF file

        Returns:
            Dictionary with extracted metadata (title, author, creation_date, etc.)
        """
        try:
            with open(file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                metadata = pdf_reader.metadata or {}

                # Extract available metadata
                result = {}
                if metadata.get('/Title'):
                    result['title'] = metadata['/Title']
                if metadata.get('/Author'):
                    result['author'] = metadata['/Author']
                if metadata.get('/CreationDate'):
                    # Parse PDF date format: "D:YYYYMMDDHHMMSS+HH'MM'"
                    pdf_date = metadata['/CreationDate']
                    if pdf_date.startswith("D:"):
                        date_str = pdf_date[2:16]  # Extract YYYYMMDDHHMMSS
                        try:
                            parsed_date = datetime.strptime(date_str, "%Y%m%d%H%M%S")
                            result['creation_date'] = parsed_date
                        except ValueError:
                            logger.warning(f"Could not parse PDF creation date: {pdf_date}")

            logger.info(f"Extracted metadata from PDF: {file_path}")
            return result
        except Exception as e:
            logger.error(f"Failed to extract metadata from PDF {file_path}: {e}")
            return {}

    @staticmethod
    def extract_metadata_from_docx(file_path: str) -> Dict[str, Any]:
        """Extract metadata from DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            Dictionary with extracted metadata (title, author, creation_date, etc.)
        """
        try:
            doc = DocxDocument(file_path)
            core_props = doc.core_properties

            result = {}
            if core_props.title:
                result['title'] = core_props.title
            if core_props.author:
                result['author'] = core_props.author
            if core_props.created:
                result['creation_date'] = core_props.created

            logger.info(f"Extracted metadata from DOCX: {file_path}")
            return result
        except Exception as e:
            logger.error(f"Failed to extract metadata from DOCX {file_path}: {e}")
            return {}

    @staticmethod
    def extract_metadata_from_html(file_path: str) -> Dict[str, Any]:
        """Extract metadata from HTML file.

        Args:
            file_path: Path to HTML file

        Returns:
            Dictionary with extracted metadata (title, author, etc.)
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as html_file:
                content = html_file.read()

            soup = BeautifulSoup(content, 'html.parser')

            result = {}

            # Extract title from <title> tag or <h1> tag
            title_tag = soup.find('title')
            if title_tag:
                result['title'] = title_tag.get_text().strip()
            else:
                h1_tag = soup.find('h1')
                if h1_tag:
                    result['title'] = h1_tag.get_text().strip()

            # Extract author from meta tags
            author_tag = soup.find('meta', attrs={'name': 'author'})
            if author_tag:
                result['author'] = author_tag.get('content', '').strip()

            # Also check for other common meta tags
            if not result.get('author'):
                author_tag = soup.find('meta', attrs={'property': 'author'})
                if author_tag:
                    result['author'] = author_tag.get('content', '').strip()

            logger.info(f"Extracted metadata from HTML: {file_path}")
            return result
        except Exception as e:
            logger.error(f"Failed to extract metadata from HTML {file_path}: {e}")
            return {}

    @staticmethod
    def extract_metadata_from_txt(file_path: str) -> Dict[str, Any]:
        """Extract basic metadata from TXT file.

        Args:
            file_path: Path to TXT file

        Returns:
            Dictionary with basic metadata (filename as title, creation date from file system)
        """
        try:
            result = {}

            # Use filename as title
            result['title'] = os.path.splitext(os.path.basename(file_path))[0]

            # Get file creation time
            stat_info = os.stat(file_path)
            result['creation_date'] = datetime.fromtimestamp(stat_info.st_ctime)

            logger.info(f"Extracted basic metadata from TXT: {file_path}")
            return result
        except Exception as e:
            logger.error(f"Failed to extract metadata from TXT {file_path}: {e}")
            return {}

    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """Extract text from PDF file.

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text content

        Per research.md Decision 1: Use PyPDF2 for PDF extraction
        """
        try:
            with open(file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            logger.info(f"Extracted {len(text)} characters from PDF: {file_path}")
            return text
        except Exception as e:
            logger.error(f"Failed to extract text from PDF {file_path}: {e}")
            raise TextExtractionError(f"PDF extraction failed: {str(e)}")

    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """Extract text from DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text content

        Per research.md Decision 1: Use python-docx for DOCX extraction
        """
        try:
            doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n".join(paragraphs)
            logger.info(f"Extracted {len(text)} characters from DOCX: {file_path}")
            return text
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX {file_path}: {e}")
            raise TextExtractionError(f"DOCX extraction failed: {str(e)}")

    @staticmethod
    def extract_from_html(file_path: str) -> str:
        """Extract text from HTML file with sanitization.

        Args:
            file_path: Path to HTML file

        Returns:
            Extracted text content with scripts/styles removed

        Per research.md Decision 1: Use BeautifulSoup4 for HTML extraction
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as html_file:
                content = html_file.read()

            soup = BeautifulSoup(content, 'html.parser')

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            # Get text and clean it up
            text = soup.get_text()

            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)

            logger.info(f"Extracted {len(text)} characters from HTML: {file_path}")
            return text
        except Exception as e:
            logger.error(f"Failed to extract text from HTML {file_path}: {e}")
            raise TextExtractionError(f"HTML extraction failed: {str(e)}")

    @staticmethod
    def extract_from_txt(file_path: str) -> str:
        """Extract text from TXT file (pass-through).

        Args:
            file_path: Path to TXT file

        Returns:
            File content as text

        Per research.md Decision 1: Simple pass-through for TXT files
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as txt_file:
                text = txt_file.read()
            logger.info(f"Extracted {len(text)} characters from TXT: {file_path}")
            return text
        except UnicodeDecodeError:
            # Try with different encoding if UTF-8 fails
            try:
                with open(file_path, 'r', encoding='latin-1') as txt_file:
                    text = txt_file.read()
                logger.info(f"Extracted {len(text)} characters from TXT (latin-1): {file_path}")
                return text
            except Exception as e:
                logger.error(f"Failed to extract text from TXT {file_path}: {e}")
                raise TextExtractionError(f"TXT extraction failed: {str(e)}")
        except Exception as e:
            logger.error(f"Failed to extract text from TXT {file_path}: {e}")
            raise TextExtractionError(f"TXT extraction failed: {str(e)}")

    @classmethod
    def extract_text(cls, file_path: str, content_type: str) -> str:
        """Extract text from file based on content type.

        Args:
            file_path: Path to file to extract text from
            content_type: Content type ('pdf', 'txt', 'docx', 'html')

        Returns:
            Extracted text content

        Per research.md Decision 1: Format-specific extractors
        """
        if content_type == 'pdf':
            return cls.extract_from_pdf(file_path)
        elif content_type == 'docx':
            return cls.extract_from_docx(file_path)
        elif content_type == 'html':
            return cls.extract_from_html(file_path)
        elif content_type == 'txt':
            return cls.extract_from_txt(file_path)
        else:
            raise TextExtractionError(f"Unsupported content type: {content_type}")

    @classmethod
    def extract_metadata(cls, file_path: str, content_type: str) -> Dict[str, Any]:
        """Extract metadata from file based on content type.

        Args:
            file_path: Path to file to extract metadata from
            content_type: Content type ('pdf', 'txt', 'docx', 'html')

        Returns:
            Dictionary with extracted metadata fields
        """
        if content_type == 'pdf':
            return cls.extract_metadata_from_pdf(file_path)
        elif content_type == 'docx':
            return cls.extract_metadata_from_docx(file_path)
        elif content_type == 'html':
            return cls.extract_metadata_from_html(file_path)
        elif content_type == 'txt':
            return cls.extract_metadata_from_txt(file_path)
        else:
            logger.warning(f"Unsupported content type for metadata extraction: {content_type}")
            return {}


class URLContentExtractor:
    """Content extraction service for URLs."""

    # Default timeout for requests (30 seconds per research.md)
    REQUEST_TIMEOUT = 30

    @classmethod
    def fetch_content_from_url(cls, url: str) -> tuple[str, str]:
        """Fetch content from URL and determine content type.

        Args:
            url: URL to fetch content from

        Returns:
            Tuple of (content, content_type) where content_type is 'html', 'pdf', etc.

        Per US5 requirements: Handle both HTML pages and PDF files from URLs
        """
        try:
            # Make request with timeout and headers
            headers = {
                'User-Agent': 'Document-Ingestion-Service/1.0'
            }
            response = requests.get(
                url,
                timeout=cls.REQUEST_TIMEOUT,
                headers=headers
            )
            response.raise_for_status()

            # Determine content type from headers
            content_type_header = response.headers.get('Content-Type', '').lower()

            # Check if it's a PDF
            if 'application/pdf' in content_type_header:
                content_type = 'pdf'
            # Check if it's HTML
            elif 'text/html' in content_type_header or 'text/' in content_type_header:
                content_type = 'html'
            # Default to HTML if content type is ambiguous
            else:
                content = response.text
                if content.strip().lower().startswith('<!doctype html') or \
                   '<html' in content[:100].lower():
                    content_type = 'html'
                else:
                    content_type = 'html'  # Default assumption

            logger.info(f"Fetched {len(response.content)} bytes from URL: {url}, type: {content_type}")
            return response.text if content_type == 'html' else response.content.decode('utf-8'), content_type

        except requests.exceptions.Timeout:
            logger.error(f"Request timeout for URL: {url}")
            raise TextExtractionError(f"Request timeout for URL: {url}")
        except requests.exceptions.RequestException as e:
            logger.error(f"Failed to fetch URL {url}: {e}")
            raise TextExtractionError(f"Failed to fetch URL: {url} - {str(e)}")
        except Exception as e:
            logger.error(f"Unexpected error fetching URL {url}: {e}")
            raise TextExtractionError(f"Unexpected error fetching URL: {url} - {str(e)}")

    @classmethod
    def save_url_content_to_temp_file(cls, url: str, content_type: str) -> str:
        """Save fetched URL content to a temporary file.

        Args:
            url: URL to fetch content from
            content_type: Expected content type ('html', 'pdf', etc.)

        Returns:
            Path to temporary file containing the content
        """
        content, detected_type = cls.fetch_content_from_url(url)

        # Create temporary file
        import tempfile
        import mimetypes

        # Determine file extension based on content type
        if content_type == 'pdf' or detected_type == 'pdf':
            ext = '.pdf'
        elif content_type == 'html' or detected_type == 'html':
            ext = '.html'
        else:
            # Default to .txt for unknown types
            ext = '.txt'

        # Create temporary file
        with tempfile.NamedTemporaryFile(mode='w', suffix=ext, delete=False) as temp_file:
            temp_file.write(content)
            temp_file_path = temp_file.name

        logger.info(f"Saved URL content to temporary file: {temp_file_path}")
        return temp_file_path


def extract_text_from_file(file_path: str, content_type: str) -> str:
    """Extract text from file using appropriate extractor.

    Args:
        file_path: Path to file to extract text from
        content_type: Content type ('pdf', 'txt', 'docx', 'html')

    Returns:
        Extracted text content
    """
    return TextExtractor.extract_text(file_path, content_type)


def extract_metadata_from_file(file_path: str, content_type: str) -> Dict[str, Any]:
    """Extract metadata from file using appropriate extractor.

    Args:
        file_path: Path to file to extract metadata from
        content_type: Content type ('pdf', 'txt', 'docx', 'html')

    Returns:
        Dictionary with extracted metadata fields
    """
    return TextExtractor.extract_metadata(file_path, content_type)


def extract_text_from_url(url: str) -> tuple[str, str]:
    """Extract text from URL content.

    Args:
        url: URL to extract content from

    Returns:
        Tuple of (extracted_text, content_type)
    """
    temp_file_path = URLContentExtractor.save_url_content_to_temp_file(url, 'html')
    try:
        content_type = 'html'  # Default assumption, could be enhanced to detect
        extracted_text = TextExtractor.extract_text(temp_file_path, content_type)
        return extracted_text, content_type
    finally:
        # Clean up temporary file
        if os.path.exists(temp_file_path):
            os.unlink(temp_file_path)